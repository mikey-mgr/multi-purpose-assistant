"""
Document generation layer.
Cover letter DOCX output (PDF resume moved to app/rendercv_renderer.py).
PDF merge utilities.
"""

import logging
import os

from app.config import settings

logger = logging.getLogger(__name__)


def ensure_output_dir(subdir: str = "") -> str:
    base = settings.OUTPUT_DIR
    path = os.path.join(base, subdir) if subdir else base
    os.makedirs(path, exist_ok=True)
    return path


def cover_letter_to_docx(text_content: str, output_path: str) -> str:
    try:
        from docx import Document
        doc = Document()
        for para in text_content.strip().split("\n\n"):
            doc.add_paragraph(para.strip())
        doc.save(output_path)
        logger.info("DOCX written to %s", output_path)
    except ImportError:
        _fallback_text(text_content, output_path.replace(".docx", ".txt"))
        logger.warning("python-docx not installed — saved as .txt instead of .docx")
    return output_path


def cover_letter_to_pdf(text_content: str, output_path: str) -> str:
    """Render cover letter text as a PDF using fpdf2."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_font("Helvetica", "", 11)

    for para in text_content.strip().split("\n\n"):
        pdf.multi_cell(0, 6, para.strip())
        pdf.ln(2)

    pdf.output(output_path)
    logger.info("Cover letter PDF written to %s", output_path)
    return output_path


def build_merged_pdf(
    required_docs: list[str],
    docs: dict,
    job_id: int,
) -> str | None:
    """Merge ALL required documents into a single PDF.

    Order: cover letter → resume → education → certification → misc docs.
    Non-PDF files (DOCX, images) are converted on a best-effort basis.
    Returns the path to the merged PDF, or None on failure.
    """
    from pypdf import PdfWriter

    # Collect file paths in desired order
    file_paths: list[str] = []

    if "cover_letter" in required_docs and docs.get("cover_letter_docx"):
        file_paths.append(docs["cover_letter_docx"])

    if "resume" in required_docs and docs.get("resume_pdf"):
        file_paths.append(docs["resume_pdf"])

    if "education_cert" in required_docs:
        file_paths.extend(docs.get("education_docs", []))

    if "certification_cert" in required_docs:
        file_paths.extend(docs.get("certification_docs", []))

    misc = docs.get("misc_docs", {})
    for doc_type in required_docs:
        if doc_type in misc:
            file_paths.append(misc[doc_type])

    if not file_paths:
        logger.warning("No files to merge")
        return None

    out_dir = ensure_output_dir("merged")
    if docs.get("resume_pdf"):
        basename = os.path.basename(docs["resume_pdf"])
    else:
        basename = f"application_{job_id}.pdf"
    merged_path = os.path.join(out_dir, basename)

    merger = PdfWriter()
    temp_files: list[str] = []

    try:
        for fp in file_paths:
            if not os.path.isfile(fp):
                logger.warning("File not found, skipping: %s", fp)
                continue

            ext = os.path.splitext(fp)[1].lower()
            if ext == ".pdf":
                merger.append(fp)
            else:
                pdf_path = _convert_to_pdf(fp, out_dir)
                if pdf_path:
                    merger.append(pdf_path)
                    if pdf_path != fp:
                        temp_files.append(pdf_path)
                else:
                    logger.warning("Could not convert %s to PDF, skipping", fp)

        if merger.pages == 0:
            logger.warning("No pages in merged PDF")
            return None

        merger.write(merged_path)
        merger.close()
        logger.info("Merged PDF written to %s (%d files)", merged_path, len(file_paths))
        return merged_path
    except Exception as e:
        logger.warning("Failed to merge PDFs: %s", e)
        return None
    finally:
        for tmp in temp_files:
            try:
                os.unlink(tmp)
            except OSError:
                pass


def _convert_to_pdf(file_path: str, output_dir: str) -> str | None:
    """Convert a non-PDF file to PDF. Returns the PDF path or None."""
    ext = os.path.splitext(file_path)[1].lower()
    basename = os.path.splitext(os.path.basename(file_path))[0]
    out_path = os.path.join(output_dir, f"{basename}.pdf")

    if ext == ".pdf":
        return file_path

    from fpdf import FPDF
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)

    if ext in (".docx", ".doc"):
        try:
            from docx import Document
            doc = Document(file_path)
            pdf.add_page()
            pdf.set_font("Helvetica", "", 11)
            for para in doc.paragraphs:
                pdf.multi_cell(0, 6, para.text)
                pdf.ln(1)
            pdf.output(out_path)
            return out_path
        except Exception as e:
            logger.warning("Failed to convert DOCX to PDF: %s", e)
            return None

    if ext in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"):
        try:
            pdf.add_page()
            pdf.image(file_path, x=10, y=10, w=190)
            pdf.output(out_path)
            return out_path
        except Exception as e:
            logger.warning("Failed to convert image to PDF: %s", e)
            return None

    if ext == ".txt":
        try:
            pdf.add_page()
            pdf.set_font("Courier", "", 10)
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                for line in f:
                    pdf.cell(0, 5, line.rstrip())
                    pdf.ln()
            pdf.output(out_path)
            return out_path
        except Exception as e:
            logger.warning("Failed to convert TXT to PDF: %s", e)
            return None

    logger.warning("Unsupported file type %s: %s", ext, file_path)
    return None


def _fallback_text(content: str, path: str) -> None:
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
